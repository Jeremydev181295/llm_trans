from docx import Document
from docx.shared import Cm
from docx2python import docx2python
import os
# import handle_docx

# Check if the file exists before attempting to remove it
def delete_paragraph(paragraph):
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

def extract_images(source_file_path):
        with docx2python(source_file_path) as docx_content:
                docx_content.save_images('./img')

def add_images(trans_file_path, result_file_path):
        # source_file_path = 'test_img.docx'
        
        # trans_file_path = 'other_test_img_result.docx'
        doc = Document(trans_file_path)
        par = doc.paragraphs
        for i in range(len(par)):
                if '----media/' in par[i].text:                        
                        start_position = par[i].text.index('media/')+6
                        end_position = len(par[i].text)-5
                        img_file_name = par[i].text[start_position:end_position]
                        img_file_path = './img/' + img_file_name
                        print(img_file_path)
                        pp = doc.paragraphs[i].insert_paragraph_before('\n')
                        pp.add_run().add_picture(img_file_path, width=Cm(15))
                        delete_paragraph(doc.paragraphs[i+1])
                        if os.path.exists(img_file_path):
                        # Remove the file
                                os.remove(img_file_path)
                                print("File removed successfully.")
                        else:
                                print("The file does not exist.")

        # result_file_path = 'demo_better.docx'      

        doc.save(result_file_path)

# add_images(source_file_path='dossier_art_my_studies.docx', trans_file_path='test_img_result.docx', result_file_path='demo_better.docx')