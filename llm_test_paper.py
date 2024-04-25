from langchain_community.llms import LlamaCpp
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from docx import Document

# Load the LlamaCpp language model, adjust GPU usage based on your hardware
llm = LlamaCpp(   
    model_path="models/llama-2-7b-chat.Q5_K_M.gguf",    
    # model_path="models/Meta-Llama-3-8B-Instruct-Q6_K.gguf",  
    # model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf", 
    n_gpu_layers=40,
    n_batch=512,  # Batch size for model processing
    verbose=False,  # Enable detailed logging for debugging
    device="cuda"
)


# file_path = 'test.docx'
file_path = 'test1.docx'

# Load the document
doc = Document(file_path)



data=[]
for paragraph in doc.paragraphs:    
    data.append(paragraph.text)
    print(paragraph.text)

trans_lang = "English"
source_lang = "French"

# Define the prompt template with a placeholder for the question
template = """
Question: {question}

Answer:
"""
# Generate translation queries for each sentence
translation_queries = [f"Translate '{sentence}' from {source_lang} to {trans_lang}. I don't need any other sentences except the translation sentences." for sentence in data]

# translation_queries = [f"Translate '{sentence}' into {trans_lang}." for sentence in data]
print(translation_queries)

# Create an LLMChain to manage interactions with the prompt and model
prompt = PromptTemplate(template=template, input_variables=["question"])
llm_chain = LLMChain(prompt=prompt, llm=llm)

print("Chatbot initialized, ready to translate...")

result_data = []
for query in translation_queries:
    answer = llm_chain.run(query)
    result_data.append(answer)
    # print(answer, '\n')

# Create a new document to store the translated data
translated_doc = Document()

# Add translated sentences to the new document
for translated_sentence in result_data:
    translated_doc.add_paragraph(translated_sentence)

# Save the translated document
translated_file_path = 'translated_test.docx'
translated_doc.save(translated_file_path)