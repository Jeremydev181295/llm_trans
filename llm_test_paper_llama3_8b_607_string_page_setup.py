from llama_cpp import Llama
from docx import Document
from docx2python import docx2python
import time
from bs4 import BeautifulSoup
from module import add_link
from module import handle_format
from module import handle_docx

llm = Llama(
    model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf",
    n_ctx=4096,      # Max tokens for in + out
    n_threads=4,     # CPU cores used
    n_gpu_layers=-1,  # Load all layers into VRAM of the GPU
    
)

source_file_path = 'test.docx'
# file_path = 'other_test1.docx'

# Load the document
# content_doc = Document(file_path)
content_doc = docx2python(source_file_path)   


trans_lang = "English"
source_lang = "French"
style = "written"

# system_prompt = "You are a helpful, smart, kind, and efficient AI assistant. You always fulfill the user's requests to the best of your ability."
# system_prompt = "You are a smart and mechanical AI language translater."
# system_prompt = f"Acts as a smart translator. Translate into {trans_lang} in {style} style. Keep all words, symbols and original style. Do not mention others. I only need translation sentence."
# system_prompt = f"Acts as a smart translator. Translate into {trans_lang} in {style} style. Keep all words, symbols and original style. Do not mention others. I only need translation sentence."
# prompt = f"Translate '{string}' from French into English. I need only translation string. Do not translate numbers and symbols and abbreviation. Keep original style."
    

# system_prompt = f"Acts as a smart translator. Translate {source_lang} sentences into {trans_lang} senteces in {style} style. Do not remove heading word. If sentence includes '----footnotes----' then translate it. I need only translation sentence."
# system_prompt = f"Translate the sentence into {trans_lang}. I need only translation sentence. Translate '----footnotes----' in the sentence. Please don't mention about others and remove extra symbols. "
def extract_content(doc):
    data_list=[]
    for sentence in doc.body[0][0][0]:
        print(sentence.strip())
        data_list.append(sentence.strip())         
    
    return data_list

def trans_with_ai(prompt, max_tokens=2048):
    """
    Function to send a prompt to the AI and return its response.
    """
    # This function sends the prompt to your AI model and fetches the response
    response = llm(prompt, max_tokens=max_tokens, temperature = 0.001, stop=["Q:", "\n"], echo=False)
    return response

def build_trans_prompt(source_lang, trans_lang, data):    
    # prompts = [f"'{sentence}'" for sentence in data if sentence.strip()]
    prompts = [f"Translate '{string.strip()}' from {source_lang} into {trans_lang}. I need only translation string. Do not translate numbers and symbols and abbreviation. Keep original style. Keep dash and colon." for string in data]

    return prompts

def main():    
    
    content_data = extract_content(content_doc) 
    content_prompts = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=content_data)
    
    
    start_time = time.time()

    result_content_data = []  
    link_target_list = []
    link_text_list = []

    # translate content of file 
    for i, prompt in enumerate(content_prompts):  
        # print(prompt)
        # prompt = system_prompt + prompt        
        if prompt:
            print(prompt)
            if "Translate '' from" in prompt:

                result_content_data.append('')
            
            else: 
                if '<a href=' in prompt:
                    # extract hyperlink text and url
                    html_string = content_data[i]

                    # Parse the HTML
                    soup = BeautifulSoup(html_string, 'html.parser')

                    # Extract the URL and text within the <a> tag
                    link_target_list.append(soup.a['href'])
                    link_text = soup.a.text
                    link_text_list.append(link_text)
                    result_content_data.append(link_text)
                else:
                    prompt = f"Q: {prompt} A: "
                    ai_response = trans_with_ai(prompt)           
                    print(ai_response)
                    result_content_data.append(ai_response['choices'][0]['text'].strip())       
        

    # translate footnote of file 

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    content_tokens_generated = sum(len(sentence.split()) for sentence in result_content_data)   


    # Create a new document to store the translated data
    translated_doc = Document()

    # Add translated content sentences to the new document
    for translated_sentence in result_content_data:
        translated_doc.add_paragraph(translated_sentence)   

    # Save the translated document
    content_translated_file_path = 'test_result.docx'
    translated_doc.save(content_translated_file_path) 

   
    # set font style
    # Load the source document
    source_doc = Document(source_file_path)
    # Load the destination document
    dest_doc = Document(content_translated_file_path)
    copy_doc = Document()
    for i, source_paragraph in enumerate(source_doc.paragraphs):
        # source_paragraph = source_doc.paragraphs[i]
        target_paragraph = dest_doc.paragraphs[i]
        copy_para = copy_doc.add_paragraph()
        handle_format.copy_paragraph_font_style(source_paragraph, target_paragraph, copy_para)
    
    add_font_style_result_path = 'test_result_font_styled.docx'
    copy_doc.save(add_font_style_result_path)
    
    # set paragraph style
    # Load the source and destination documents
    source_doc = Document(source_file_path)
    # print(source_doc.styles.)
    target_doc = Document(add_font_style_result_path)
    handle_format.copy_paragraph_style(source_doc, target_doc)

    add_para_style_result_path = 'test_result_para_styled.docx'
    target_doc.save(add_para_style_result_path)

    # add hyperlinks
    document = Document(add_para_style_result_path)
    for para in document.paragraphs:
        for i, link_text in enumerate(link_text_list):
            if para.text == link_text:
                para.text = ''
                add_link.add_hyperlink(para, source_file_path, link_text, link_target_list[i])
    
    hyperlink_result_path = 'test_result_hyperlink.docx'
    document.save(hyperlink_result_path)

    
    handle_docx.copy_page_setup(source_file_path=source_file_path, target_file_path=hyperlink_result_path)
    
    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    
   
# To start the chat, call the main_chat function
main()