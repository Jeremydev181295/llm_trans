from llama_cpp import Llama
from spire.doc import *
from spire.doc.common import *
import time
from module import handle_docx
import docx

llm = Llama(
    model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf",
    n_ctx=4096,      # Max tokens for in + out
    n_threads=4,     # CPU cores used
    n_gpu_layers=-1,  # Load all layers into VRAM of the GPU
    
)

source_file_path = 'test1.docx'

trans_lang = "English"
source_lang = "French"
style = "written"

def trans_with_ai(prompt, max_tokens=2048):
    """
    Function to send a prompt to the AI and return its response.
    """
    # This function sends the prompt to your AI model and fetches the response
    response = llm(prompt, max_tokens=max_tokens, temperature = 0.001, stop=["Q:", "\n"], echo=False)
    return response

def build_trans_prompt(source_lang, trans_lang, string):    
    # prompts = [f"'{sentence}'" for sentence in data if sentence.strip()]
    prompt = f"Translate '{string.strip()}' from {source_lang} into {trans_lang}. I need only translation. Do not translate numbers, symbols and abbreviations. Keep original style. Keep dash and colon. "

    return prompt

def main():    
    
    document = Document()
    document.LoadFromFile(source_file_path)
    
    start_time = time.time()
    
    ############################################################################################################
    # translate content
    for index in range(len(document.Sections)):
        section = document.Sections[index]
        for i in range(len(section.Paragraphs)):
            string = section.Paragraphs[i].Text
            if not string =='':
                if not 'https://' in string:
                    prompt = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, string=string)
                    prompt = f"Q: {prompt} A: "
                    ai_response = trans_with_ai(prompt)           
                    print(ai_response)
                    section.Paragraphs[i].Text = ai_response['choices'][0]['text'].strip()

    content_translated_file_path = 'content_translated_result.docx'
    document.SaveToFile(content_translated_file_path, FileFormat.Docx)

    ###########################################################################################################
    # remove unnecessary first paragraph
    revise_doc = docx.Document(content_translated_file_path)
    if revise_doc.paragraphs[0].text == "Evaluation Warning: The document was created with Spire.Doc for Python.":
            handle_docx.delete_paragraph(revise_doc.paragraphs[0])
    revise_doc.save(content_translated_file_path)

    ###########################################################################################################
    
    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    
   
# To start the chat, call the main_chat function
main()