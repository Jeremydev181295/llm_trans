from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



from docx import Document

# remove paragraphs using docx package
def remove_empty_paragraphs(doc):
    for para in doc.paragraphs:
        if not para.text.strip():
            p_element = para._element
            p_element.getparent().remove(p_element)

def extract_header(doc):
    header_data_list = []
    for header_run in doc.header_runs:
        if header_run[0][0][0]:
            for data in header_run[0][0]:
                item_string = ''
                for item in data:
                    item_string += item
                if item_string:
                    print(item_string)   
                    header_data_list.append(item_string)
    return header_data_list  
    
# extract_content using docx2python
def extract_content(doc):
    data_list=[]
    for sentence in doc.body[0][0][0]:
        print(sentence.strip())
        data_list.append(sentence.strip())         
    
    return data_list
def extract_footnote_pos_string(doc, footnote_string):
    pos = 0
    for sentence in doc.body[0][0][0]:
        string = sentence.strip()
        p = string.find(footnote_string)
        if p != -1:
            pos = p
            break
    return string[pos-30:pos]

def extract_footnote_para_strings(doc, index_list):
    footnote_para_strings = []
    for index in index_list:
        footnote_para_strings.append(doc.body[0][0][0][index].strip())
    return footnote_para_strings

# extract_footnote using docx2python
def extract_footnote(doc):
    footnote_list = []
    for footnote in doc.footnotes_runs[0][0]:
        for specific in footnote:
            for line in specific:
                split_lines = line.split("\t")
                if "footnote" not in split_lines[0]:                    
                    footnote_list.append(split_lines[0].strip())
    return footnote_list

# # extract_content using Document from docx
# def extract_content(doc):
#     data_list=[]
#     for paragraph in doc.paragraphs:    
#         data_list.append(paragraph.text.strip())
#     return data_list

# handle paragraph using Document from docx-python package
def remove_string_from_paragraph(doc, target_string):
    for paragraph in doc.paragraphs:            
        if target_string in paragraph.text:
            print(paragraph.text)
            paragraph.text = paragraph.text.replace(target_string, '')

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_page_setup(source_file_path, target_file_path):
    source_doc = Document(source_file_path)
    target_doc = Document(target_file_path)
    source_doc_section = source_doc.sections[0]
    target_doc_section = target_doc.sections[0]

    target_doc_section.page_height = source_doc_section.page_height
    target_doc_section.page_width = source_doc_section.page_width
    target_doc_section.left_margin = source_doc_section.left_margin
    target_doc_section.right_margin = source_doc_section.right_margin
    target_doc_section.top_margin = source_doc_section.top_margin
    target_doc_section.bottom_margin = source_doc_section.bottom_margin
    target_doc_section.header_distance = source_doc_section.header_distance
    target_doc_section.footer_distance = source_doc_section.footer_distance
    
    target_doc.save(target_file_path)

def set_space_para_same_style(source_file_path, target_file_path):
    source_doc = Document(source_file_path)
    target_doc = Document(target_file_path)
    for i, p in enumerate(source_doc.paragraphs):
        p_element = p._element
        if i < len(target_doc.paragraphs):
            target_p_element = target_doc.paragraphs[i]._element
            contextual_spacings_source = p_element.xpath(r'w:pPr/w:contextualSpacing')
            contextual_spacings_target = target_p_element.xpath(r'w:pPr/w:contextualSpacing')
            if contextual_spacings_source:
                add_contextual_spacing(target_doc.paragraphs[i])
            else:
                if contextual_spacings_target:
                    cspacing = contextual_spacings_target[0]
                    cspacing.getparent().remove(cspacing)

    target_doc.save(target_file_path)
    

def add_contextual_spacing(paragraph):
    pPr = paragraph._element.pPr
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    
    # Check if contextualSpacing already exists
    existing_contextual_spacing = pPr.find(qn('w:contextualSpacing'))
    if existing_contextual_spacing is None:
        # Create the contextualSpacing element
        contextual_spacing = OxmlElement('w:contextualSpacing')
        pPr.append(contextual_spacing)


